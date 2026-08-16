"""多模板排版输出引擎。

- render_html : 渲染 HTML（预览与 PDF 底座）
- render_pdf  : 通过 Playwright/Chromium 导出 PDF
- render_docx : 通过 python-docx 导出 Word
"""

import html
import json
import re
from io import BytesIO
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = BASE_DIR / "templates"

MASK = "***"


def _e(v) -> str:
    """转义并转字符串，None 视为空。"""
    if v is None:
        return ""
    return html.escape(str(v))


def _mask(value: str, sensitive: bool) -> str:
    return MASK if (sensitive and value) else _e(value)


def _apply_sensitive(resume: dict, hide_sensitive: bool) -> dict:
    """复制简历并打码敏感字段，不修改原数据。"""
    import copy

    if not hide_sensitive:
        return resume
    out = copy.deepcopy(resume)
    marked = set(out.get("sensitive_fields", []))
    basic = out.get("basic", {})
    for key in ("name", "phone", "email", "city"):
        if key in marked and basic.get(key):
            basic[key] = MASK
    if "company" in marked:
        for item in out.get("work", []):
            if item.get("company"):
                item["company"] = MASK
    return out


def list_templates() -> list[dict]:
    """返回模板元信息列表。"""
    result = []
    for path in sorted(TEMPLATE_DIR.glob("*.json")):
        with open(path, encoding="utf-8") as f:
            tpl = json.load(f)
        result.append({"id": tpl["id"], "name": tpl.get("name", tpl["id"]), "layout": tpl.get("layout")})
    return result


def load_template(template_id: str) -> dict:
    path = TEMPLATE_DIR / f"{template_id}.json"
    if not path.exists():
        raise ValueError(f"模板不存在: {template_id}")
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def _sections_html(resume: dict, tpl: dict) -> str:
    basic = resume.get("basic", {})
    sections = {s["key"]: s for s in tpl.get("sections", [])}
    blocks = []

    edu = resume.get("education", [])
    if "education" in sections and edu:
        items = "".join(
            f'<div class="entry"><div class="entry-head"><span class="entry-title">{_e(i.get("school"))}</span>'
            f'<span class="entry-sub">{_e(i.get("degree"))} · {_e(i.get("major"))}</span>'
            f'<span class="entry-date">{_e(i.get("graduation"))}</span></div>'
            f'<div class="entry-note">{_e(i.get("note"))}</div></div>'
            for i in edu
        )
        blocks.append((sections["education"]["order"], '<section><h2>教育背景</h2>' + items + "</section>"))

    work = resume.get("work", [])
    if "work" in sections and work:
        items = "".join(
            '<div class="entry"><div class="entry-head"><span class="entry-title">%s</span>'
            '<span class="entry-sub">%s</span><span class="entry-date">%s - %s</span></div><ul>%s</ul></div>'
            % (
                _e(w.get("company")),
                _e(w.get("position")),
                _e(w.get("start")),
                _e(w.get("end")),
                "".join(f"<li>{_e(r)}</li>" for r in w.get("responsibilities", [])),
            )
            for w in work
        )
        blocks.append((sections["work"]["order"], '<section><h2>工作经历</h2>' + items + "</section>"))

    projects = resume.get("projects", [])
    if "projects" in sections and projects:
        items = "".join(
            '<div class="entry"><div class="entry-head"><span class="entry-title">%s</span>'
            '<span class="entry-sub">%s</span><span class="entry-date">%s - %s</span></div><ul>%s</ul>'
            '<div class="entry-note">%s</div></div>'
            % (
                _e(p.get("name")),
                _e(p.get("role")),
                _e(p.get("start")),
                _e(p.get("end")),
                "".join(f"<li>{_e(d)}</li>" for d in p.get("description", [])),
                "".join(f"<li class='achv'>成果：{_e(a)}</li>" for a in p.get("achievements", []))
                if p.get("achievements")
                else "",
            )
            for p in projects
        )
        blocks.append((sections["projects"]["order"], '<section><h2>项目经验</h2>' + items + "</section>"))

    if "skills" in sections:
        skills = resume.get("skills", {})
        parts = []
        if skills.get("skills"):
            parts.append("技能：" + " / ".join(_e(s) for s in skills["skills"]))
        if skills.get("certificates"):
            parts.append("证书：" + " / ".join(_e(s) for s in skills["certificates"]))
        if skills.get("languages"):
            parts.append("语言：" + " / ".join(_e(s) for s in skills["languages"]))
        if parts:
            blocks.append((sections["skills"]["order"], "<section><h2>技能证书</h2><p>" + "<br>".join(parts) + "</p></section>"))

    if "self_eval" in sections and resume.get("self_evaluation"):
        blocks.append(
            (sections["self_eval"]["order"], "<section><h2>自我评价</h2><p>" + _e(resume["self_evaluation"]) + "</p></section>")
        )

    blocks.sort(key=lambda x: x[0])
    return "".join(b[1] for b in blocks)


def render_html(resume: dict, tpl: dict, hide_sensitive: bool = False) -> str:
    """渲染模板为 HTML 字符串。"""
    resume = _apply_sensitive(resume, hide_sensitive)
    basic = resume.get("basic", {})
    style = tpl.get("style", {})

    title_style = "border-bottom:2px solid %s; padding-bottom:4px;" % style.get("primary_color", "#333") if style.get(
        "section_title_style"
    ) == "line" else "background:%s; color:#fff; padding:4px 10px;" % style.get("primary_color", "#333")

    if tpl.get("layout") == "two_column_header":
        header = (
            f'<div class="header two"><div class="h-main"><h1>{_e(basic.get("name"))}</h1>'
            f'<div class="h-sub">{_e(basic.get("job_intention"))}</div></div>'
            f'<div class="h-side">{_e(basic.get("city"))}<br>{_e(basic.get("phone"))}<br>{_e(basic.get("email"))}</div></div>'
        )
    else:
        header = (
            f'<div class="header single"><h1>{_e(basic.get("name"))}</h1>'
            f'<div class="h-sub">{_e(basic.get("job_intention"))}'
            + (" · " if basic.get("city") else "")
            + f'{_e(basic.get("city"))}</div>'
            f'<div class="h-line">{_e(basic.get("phone"))}'
            + (" ｜ " if basic.get("phone") and basic.get("email") else " ")
            + f'{_e(basic.get("email"))}</div></div>'
        )

    return f"""<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="utf-8">
<style>
  * {{ margin:0; padding:0; box-sizing:border-box; }}
  body {{ font-family:{style.get("font_family", "sans-serif")}; font-size:{style.get("font_size", "14px")}; color:#222; padding:{style.get("page_padding", "36px")}; background:#fff; }}
  h1 {{ font-size:26px; color:{style.get("primary_color", "#222")}; }}
  .header {{ margin-bottom:18px; }}
  .header.two {{ display:flex; justify-content:space-between; align-items:flex-end; }}
  .h-sub {{ color:#666; margin-top:4px; }}
  .h-line {{ color:#666; margin-top:4px; }}
  .h-side {{ text-align:right; color:#555; line-height:1.7; }}
  section {{ margin-bottom:16px; }}
  h2 {{ font-size:16px; color:{style.get("primary_color", "#222")}; margin-bottom:8px; {title_style} }}
  .entry {{ margin-bottom:10px; }}
  .entry-head {{ display:flex; justify-content:space-between; align-items:baseline; }}
  .entry-title {{ font-weight:600; }}
  .entry-sub {{ color:#555; }}
  .entry-date {{ color:#888; font-size:13px; }}
  .entry-note {{ color:#555; margin-top:2px; }}
  ul {{ margin:4px 0 0 18px; }}
  li {{ line-height:1.6; }}
  .achv {{ color:#2563eb; }}
</style></head>
<body>
{header}
{_sections_html(resume, tpl)}
</body></html>"""


def render_pdf(resume: dict, tpl: dict, hide_sensitive: bool = False) -> bytes:
    """通过 Playwright + Chromium 将 HTML 渲染为 PDF。需先安装：playwright install chromium"""
    from playwright.sync_api import sync_playwright

    html_text = render_html(resume, tpl, hide_sensitive)
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.set_content(html_text, wait_until="networkidle")
        pdf = page.pdf(format="A4", print_background=True, margin={"top": "0", "bottom": "0", "left": "0", "right": "0"})
        browser.close()
    return pdf


def render_docx(resume: dict, tpl: dict, hide_sensitive: bool = False) -> bytes:
    """使用 python-docx 生成 Word 文档。"""
    import docx
    from docx.shared import Pt, RGBColor

    resume = _apply_sensitive(resume, hide_sensitive)
    basic = resume.get("basic", {})
    style_cfg = tpl.get("style", {})
    primary = RGBColor(*[int(style_cfg.get("primary_color", "#222")[i:i + 2], 16) for i in (1, 3, 5)])

    doc = docx.Document()
    h1 = doc.add_heading(_e(basic.get("name")) or "姓名", level=1)
    for run in h1.runs:
        run.font.color.rgb = primary
    line = doc.add_paragraph(_e(basic.get("job_intention")))
    line.add_run("  ｜  " + _e(basic.get("phone"))).font.size = Pt(11)
    line.add_run("  ｜  " + _e(basic.get("email"))).font.size = Pt(11)

    def _section(title: str, content: str):
        head = doc.add_heading(title, level=2)
        for run in head.runs:
            run.font.color.rgb = primary
        for part in content.splitlines():
            if part.strip():
                p = doc.add_paragraph(part)
                p.paragraph_format.space_after = Pt(2)

    def _entries(items, head_keys, body_key=None, body_extra=None):
        for it in items:
            head = " | ".join(_e(it.get(k)) for k in head_keys if it.get(k))
            _section(head, "")
            if body_key:
                for r in it.get(body_key, []):
                    p = doc.add_paragraph(r, style="List Bullet")
                    p.paragraph_format.space_after = Pt(1)
            if body_extra:
                extra = "；".join(_e(v) for v in it.get(body_extra, []))
                if extra:
                    doc.add_paragraph("成果：" + extra)

    _entries(resume.get("education", []), ["school", "degree", "major", "graduation"], body_extra="note")
    _entries(resume.get("work", []), ["company", "position", "start", "end"], body_key="responsibilities")
    _entries(resume.get("projects", []), ["name", "role", "start", "end"], body_key="description", body_extra="achievements")

    skills = resume.get("skills", {})
    if any(skills.get(k) for k in ("skills", "certificates", "languages")):
        lines = []
        if skills.get("skills"):
            lines.append("技能：" + " / ".join(_e(s) for s in skills["skills"]))
        if skills.get("certificates"):
            lines.append("证书：" + " / ".join(_e(s) for s in skills["certificates"]))
        if skills.get("languages"):
            lines.append("语言：" + " / ".join(_e(s) for s in skills["languages"]))
        _section("技能证书", "\n".join(lines))

    if resume.get("self_evaluation"):
        _section("自我评价", _e(resume["self_evaluation"]))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()